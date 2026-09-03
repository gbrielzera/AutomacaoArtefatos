import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import copy
from pygments import lex
from pygments.lexers import PythonLexer
from pygments.token import Token
import os
import re
import sys
import json
import logging
from datetime import datetime

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    DND_DISPONIVEL = True
except ImportError:
    # tkinterdnd2 é opcional: se não estiver instalado, o app cai para tk.Tk() normal
    # e simplesmente não oferece arrastar-e-soltar (não deve travar por isso).
    DND_DISPONIVEL = False


def get_base_dir():
    """Retorna a pasta do executável (.exe) ou do script .py, para localizar a pasta de logs."""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def limpar_nome_arquivo(nome):
    """Remove caracteres que o Windows não aceita em nomes de arquivo."""
    for c in '<>:"/\\|?*':
        nome = nome.replace(c, "_")
    return nome.strip() or "Supravizio"


class BufferLogHandler(logging.Handler):
    """
    Acumula as mensagens em memória para que cada artefato gerado receba o seu próprio
    arquivo de log, contendo tudo desde a geração anterior (inclusive a seleção do XML).
    """

    def __init__(self):
        super().__init__()
        self.linhas = []

    def emit(self, record):
        self.linhas.append(self.format(record))

    def drenar(self):
        linhas, self.linhas = self.linhas, []
        return linhas


def setup_logger():
    """Configura o logger da aplicação com um buffer em memória."""
    logger = logging.getLogger("SupravizioDocApp")
    logger.setLevel(logging.DEBUG)
    logger.handlers.clear()

    handler = BufferLogHandler()
    handler.setLevel(logging.DEBUG)
    handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
    logger.addHandler(handler)

    return logger, handler


def escrever_arquivo_log(linhas, sufixo):
    """
    Grava uma lista de linhas já formatadas em <base_dir>/logs/log_{sufixo}_{timestamp}.log.
    Usada tanto para o log de cada geração de artefato quanto para o log de uma
    exceção não tratada (sufixo "crash"), evitando duas implementações divergentes.
    """
    if not linhas:
        return ""
    try:
        logs_dir = os.path.join(get_base_dir(), "logs")
        os.makedirs(logs_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_path = os.path.join(logs_dir, f"log_{limpar_nome_arquivo(sufixo)}_{timestamp}.log")
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("\n".join(linhas) + "\n")
        return log_path
    except Exception as e:
        # Último recurso: se nem o log conseguiu ser gravado, não há mais para onde
        # reportar além do console.
        print(f"Não foi possível gravar o arquivo de log: {e}")
        return ""


def criar_manipulador_excecoes(logger, log_handler):
    """
    Cria um manipulador de exceções não tratadas, usável tanto como `sys.excepthook`
    (exceções que escapam de tudo, ex.: durante __init__ antes da GUI existir) quanto
    como `root.report_callback_exception` do Tkinter (exceções dentro de callbacks de
    botão/diálogo durante o mainloop() — o Tk NÃO propaga essas para sys.excepthook,
    por padrão só imprime no console e segue em frente sem deixar rastro).
    Ambas têm a mesma assinatura (exc_type, exc_value, exc_tb), então uma função serve
    para as duas. Sempre loga e grava um log de crash antes de avisar o usuário.
    """
    def manipulador(exc_type, exc_value, exc_tb):
        if issubclass(exc_type, KeyboardInterrupt):
            sys.__excepthook__(exc_type, exc_value, exc_tb)
            return

        logger.critical("Exceção não tratada", exc_info=(exc_type, exc_value, exc_tb))
        log_path = escrever_arquivo_log(log_handler.drenar(), "crash")
        try:
            messagebox.showerror(
                "Erro inesperado",
                f"Ocorreu um erro inesperado e a operação foi interrompida:\n{exc_value}"
                f"\n\nDetalhes em:\n{log_path}"
            )
        except Exception:
            pass  # se nem a caixa de diálogo funcionar, o log já foi gravado

    return manipulador


def dividir_caminhos_dnd(data):
    """
    Faz o parsing de event.data de um drop do tkinterdnd2 SEM usar tk.splitlist.
    tk.splitlist processa `data` como lista Tcl, onde `\\` é caractere de escape —
    isso corrompe caminhos do Windows sem chaves (ex.: "C:\\pasta\\arquivo.xml" vira
    "C:pastarquivo.xml", perdendo as barras). Aqui, caminhos entre chaves {...} são
    tratados como um único item (podem conter espaço); fora delas, a separação é por
    espaço — sem qualquer interpretação de barra invertida.
    """
    caminhos = []
    atual = ""
    dentro_chaves = False
    for ch in data:
        if ch == "{":
            dentro_chaves = True
            continue
        if ch == "}":
            dentro_chaves = False
            continue
        if ch == " " and not dentro_chaves:
            if atual:
                caminhos.append(atual)
                atual = ""
            continue
        atual += ch
    if atual:
        caminhos.append(atual)
    return caminhos


def get_config_path():
    return os.path.join(get_base_dir(), "config", "settings.json")


def carregar_config(logger=None):
    """
    Lê o histórico salvo da última execução (Macroprocesso, Processo, caminhos).
    Nunca lança: um arquivo ausente é o caso normal de primeira execução (sem aviso);
    um arquivo corrompido/inacessível vira aviso no log e o app segue com config vazia.
    """
    caminho = get_config_path()
    if not os.path.exists(caminho):
        return {}
    try:
        with open(caminho, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        if logger:
            logger.warning(f"Não foi possível ler o arquivo de configurações '{caminho}': {e}")
        return {}


def salvar_config(dados, logger=None):
    """Grava o histórico em disco. Nunca lança: uma falha aqui não pode derrubar a geração."""
    caminho = get_config_path()
    try:
        os.makedirs(os.path.dirname(caminho), exist_ok=True)
        with open(caminho, "w", encoding="utf-8") as f:
            json.dump(dados, f, ensure_ascii=False, indent=2)
    except Exception as e:
        if logger:
            logger.warning(f"Não foi possível salvar o arquivo de configurações '{caminho}': {e}")


# ==========================================================
# LÓGICA DE EXTRAÇÃO DE DADOS (funções de módulo, testáveis sem GUI)
# ==========================================================
def get_texto(node, tag, default=""):
    child = node.find(tag)
    if child is not None and child.text:
        return child.text.strip()
    return default


def extrair_propriedades_campos(root):
    propriedades = {}

    for prop in root.findall(".//CustomProperty"):
        nome = get_texto(prop, "Name")
        tabela = get_texto(prop, "TableName")

        if not nome or not tabela:
            continue

        if nome not in propriedades:
            propriedades[nome] = {
                "rotulo": get_texto(prop, "Text") or get_texto(prop, "Description") or nome,
                "descricao": get_texto(prop, "Description") or get_texto(prop, "Text") or nome,
                "tipo": get_texto(prop, "Type", "String"),
                "tabela": tabela,
                "coluna": get_texto(prop, "TableColumn"),
                # Controle da tela (DropDownList, ListBox, CheckBox, TextBox...) e, para os
                # de lista, as opções disponíveis separadas por ';'. Alimentam a coluna
                # "Lista de Itens" da tabela de Campos alterados.
                "controle": get_texto(prop, "Control"),
                "lista_itens": get_texto(prop, "ListItems"),
            }

    return propriedades


def separar_lista_itens(texto):
    """Converte o conteúdo de <ListItems> ('A;B;C') em lista, ignorando itens vazios."""
    if not texto:
        return []
    return [item.strip() for item in texto.split(";") if item.strip()]


def ler_xml_root(xml_path):
    try:
        return ET.parse(xml_path).getroot()
    except ET.ParseError:
        with open(xml_path, 'r', encoding='utf-8-sig') as file:
            xml_content = file.read().strip()
        if not xml_content:
            raise Exception("O arquivo XML selecionado está vazio.")
        return ET.fromstring(xml_content)


def extrair_processo_do_nome_arquivo(caminho_xml):
    """
    O Supravizio exporta o XML com o padrão:
        "{Processo}_Versão_{N}_{NomeSubProcesso}"
    O Processo pode conter hífens (ex.: "Financeiro - Serviços Gerais"), por isso
    tudo que vem antes de "_Versão_{N}_" é considerado o Processo inteiro.
    O Macroprocesso NÃO consta nem no nome do arquivo nem no XML.
    Retorna (processo, subprocesso_do_nome); cada item pode ser None.
    """
    nome_base = os.path.splitext(os.path.basename(caminho_xml))[0]
    nome_normalizado = re.sub(r"\s+", " ", nome_base.replace("_", " ")).strip()

    match = re.match(
        r"^(?P<processo>.+?)\s+Vers[aã]o\s+\d+\s*(?P<sub>.*)$",
        nome_normalizado,
        flags=re.IGNORECASE
    )
    if not match:
        return None, None

    processo = match.group("processo").strip(" -")
    sub = match.group("sub").strip()
    return (processo or None), (sub or None)


def extrair_macroprocesso_do_xml(root):
    """
    Procura o Macroprocesso na árvore do XML. O Supravizio não exporta o nome do
    macroprocesso no XML de fluxo (apenas a classe vazia
    'Venki.Supravizio.Processo.MacroProcesso'), então na prática isto quase sempre
    retorna None e o campo precisa ser preenchido manualmente.
    """
    for tag in ("MacroProcesso", "Macroprocesso", "NomeMacroProcesso", "ClasseMacroProcesso"):
        for node in root.iter(tag):
            if node.text and node.text.strip():
                return node.text.strip()
            desc = node.find("Descricao")
            if desc is not None and desc.text and desc.text.strip():
                return desc.text.strip()
    return None


def extrair_dados_xml(xml_path, logger):
    logger.info(f"Iniciando extração de dados do XML: {xml_path}")
    root = ler_xml_root(xml_path)

    dados = {"nome_fluxo": "", "servicos": [], "campos": [], "anexos": [], "scripts": []}
    propriedades_campos = extrair_propriedades_campos(root)

    node_nome = root.find(".//NomeSubProcesso")
    if node_nome is not None and node_nome.text:
        dados["nome_fluxo"] = node_nome.text.strip()
    else:
        logger.warning("Não foi possível localizar a tag <NomeSubProcesso> no XML.")

    # Dedup por (local, código): evita duplicar o mesmo script quando a MESMA atividade
    # aparece repetida no diagrama (ex: um LinkInicial compartilhado entre várias páginas),
    # mas preserva scripts com código idêntico que estejam em locais distintos (ex: o mesmo
    # preenchimento de campo replicado no Evento Inicial e no Link Inicial).
    scripts_vistos = set()

    def registrar_script(local, sc_code):
        chave = (local, sc_code)
        if chave in scripts_vistos:
            logger.debug(f"Script duplicado ignorado (mesma atividade repetida no diagrama): {local}")
            return
        scripts_vistos.add(chave)
        dados["scripts"].append({"local": local, "codigo": sc_code})

    servicos_vistos = set()

    def coletar_servicos(restricoes):
        """Lê pares (tipo, serviço) de nós <RestricaoServico>, ignorando entradas vazias."""
        for rest in restricoes:
            srv = rest.find("Servico")
            tipo_txt = (
                get_texto(rest, "ClasseServico/Descricao")
                or (get_texto(srv, "ClasseServico/Descricao") if srv is not None else "")
            )
            nome_txt = get_texto(srv, "Descricao") if srv is not None else ""

            if not tipo_txt and not nome_txt:
                logger.warning(
                    "Encontrada uma <RestricaoServico> sem Tipo e sem Descrição do serviço; "
                    "entrada ignorada (o serviço não veio expandido no XML)."
                )
                continue

            chave = f"{tipo_txt} - {nome_txt}"
            if chave not in servicos_vistos:
                servicos_vistos.add(chave)
                dados["servicos"].append({"tipo": tipo_txt, "nome": nome_txt})

    # Fonte autoritativa: os serviços do PRÓPRIO subprocesso (ClasseSubProcesso).
    # Isto evita capturar serviços de subprocessos secundários apenas engatilhados,
    # e funciona mesmo em fluxos sem LinkInicial.
    coletar_servicos(root.findall(".//SubProcesso/ClasseSubProcesso/RestricoesServicos/RestricaoServico"))

    if not dados["servicos"]:
        logger.warning(
            "Nenhum serviço em <SubProcesso/ClasseSubProcesso>; tentando o LinkInicial como alternativa."
        )
        for link_inicial in root.findall(".//Atividade[Tipo='LinkInicial']"):
            for alvo in link_inicial.findall(".//ClasseAlvo"):
                coletar_servicos(alvo.findall(".//RestricoesServicos/RestricaoServico"))

    mapa_scripts = {
        "ScriptModificado": "modificado",
        "ScriptValidacao": "de validação",
        "ScriptFormCarregado": "de formulário carregado",
        "ScriptInicio": "de início",
        "ScriptFim": "de fim",
        "ScriptVolta": "de volta",
        "ScriptEvento": "de evento",
    }

    # Scripts que ficam aninhados sob outro nó da atividade, e não como filho direto.
    mapa_scripts_aninhados = {
        "PapelResponsavel/PapelClasseNegocio/ScriptSelecaoAtores": "de seleção de atores (papel responsável)",
        "PapelDestinatario/PapelClasseNegocio/ScriptSelecaoAtores": "de seleção de atores (papel destinatário)",
    }

    campos_vistos = set()
    anexos_vistos = set()

    for figura in root.findall(".//Figura"):
        atividade = figura.find("Atividade")

        if atividade is not None:
            tag_tipo_atv = atividade.find("Tipo")
            tipo_atv = tag_tipo_atv.text.strip() if (tag_tipo_atv is not None and tag_tipo_atv.text) else "Atividade"

            # Nome real da tarefa (ex.: "Anexar comprovante de pagamento"), quando existir.
            # Tipos sem rótulo próprio (LinkInicial, Gateway, etc.) caem no fallback do Tipo.
            tag_desc_atv = atividade.find("Descricao")
            descricao_atv = tag_desc_atv.text.strip() if (tag_desc_atv is not None and tag_desc_atv.text) else ""
            if descricao_atv and descricao_atv != tipo_atv:
                nome_atv = f"{descricao_atv} ({tipo_atv})"
            else:
                nome_atv = tipo_atv

            # As operações do LinkInicial (evento de inicialização do fluxo) ficam em
            # <Figura><OperacaoAtividade>, IRMÃO de <Atividade> e não dentro dela.
            # Ler só atividade//OperacaoAtividade perdia esses scripts e campos.
            operacoes = atividade.findall(".//OperacaoAtividade") + figura.findall("OperacaoAtividade")

            for op in operacoes:
                for campo in op.findall(".//CampoPreenchimento"):
                    nome_custom = campo.find("./NomeCustomizado")
                    rotulo = campo.find("./Rotulo")

                    # Extrai a Tabela da Tag 'FormaEdicaoWeb' no nível do Campo (Corrigido)
                    tabela_nome = "Formulário"
                    form_edicao = campo.find("./FormaEdicaoWeb")
                    if form_edicao is not None and form_edicao.text:
                        if form_edicao.text.strip().lower() != "formulario":
                            tabela_nome = form_edicao.text.strip()

                    if nome_custom is not None and nome_custom.text:
                        nome_c = nome_custom.text.strip()
                        prop_campo = propriedades_campos.get(nome_c, {})
                        rot_txt = rotulo.text.strip() if (rotulo is not None and rotulo.text) else prop_campo.get("rotulo", nome_c)
                        desc_txt = prop_campo.get("descricao", rot_txt)
                        tipo_txt = prop_campo.get("tipo", "String")
                        tabela_nome = prop_campo.get("tabela", tabela_nome)

                        if nome_c and nome_c not in campos_vistos:
                            dados["campos"].append({
                                "nome": nome_c,
                                "rotulo": rot_txt,
                                "descricao": desc_txt,
                                "tipo": tipo_txt,
                                "tabela": tabela_nome,
                                "controle": prop_campo.get("controle", ""),
                                "lista_itens": prop_campo.get("lista_itens", ""),
                            })
                            campos_vistos.add(nome_c)

                        for tag_xml, nome_amigavel in mapa_scripts.items():
                            s_node = campo.find(tag_xml)
                            if s_node is not None and s_node.text and s_node.text.strip():
                                sc_code = s_node.text.strip()
                                registrar_script(
                                    f"Atividade: {nome_atv} | Script {nome_amigavel} no campo: {nome_c}",
                                    sc_code
                                )

                for tag_xml, nome_amigavel in mapa_scripts.items():
                    s_node = op.find(tag_xml)
                    if s_node is not None and s_node.text and s_node.text.strip():
                        sc_code = s_node.text.strip()
                        registrar_script(
                            f"Atividade: {nome_atv} | Script {nome_amigavel} da operação",
                            sc_code
                        )

            for tag_xml, nome_amigavel in list(mapa_scripts.items()) + list(mapa_scripts_aninhados.items()):
                s_node = atividade.find(tag_xml)
                if s_node is not None and s_node.text and s_node.text.strip():
                    sc_code = s_node.text.strip()
                    registrar_script(
                        f"Atividade: {nome_atv} | Script {nome_amigavel} da atividade",
                        sc_code
                    )

            for vi in atividade.findall(".//ValoresInputs/ValorInput"):
                expr_in = vi.find("ExpressaoValor")
                if expr_in is not None and expr_in.text and expr_in.text.strip():
                    nome_input = get_texto(vi, "CustomProperty/Name") or get_texto(vi, "Nome") or "input"
                    registrar_script(
                        f"Atividade: {nome_atv} | Expressão de valor do input: {nome_input}",
                        expr_in.text.strip()
                    )

            for escopo in atividade.findall(".//EscopoClasseAnexo/ClasseConfiguracao"):
                desc = escopo.find("./Descricao")
                sigla = escopo.find("./Sigla")
                if desc is not None and sigla is not None:
                    d_txt = desc.text.strip() if desc.text else ""
                    s_txt = sigla.text.strip() if sigla.text else ""
                    if d_txt and s_txt and s_txt not in anexos_vistos:
                        dados["anexos"].append({"nome": d_txt, "sigla": s_txt})
                        anexos_vistos.add(s_txt)

        gateway = figura.find("Gateway")
        if gateway is not None:
            tag_desc_gw = gateway.find("Descricao")
            desc_gw = tag_desc_gw.text.strip() if (tag_desc_gw is not None and tag_desc_gw.text) else "Gateway"
            expr = gateway.find("ExpressaoComparacaoDecision")

            if expr is not None and expr.text and expr.text.strip():
                sc_code = expr.text.strip()
                registrar_script(
                    f"Gateway de Decisão: {desc_gw} | Expressão de Decisão",
                    sc_code
                )

    logger.info(
        f"Extração concluída. nome_fluxo='{dados['nome_fluxo']}' | "
        f"servicos={len(dados['servicos'])} | campos={len(dados['campos'])} | "
        f"anexos={len(dados['anexos'])} | scripts={len(dados['scripts'])}"
    )
    if not dados["servicos"]:
        logger.warning("Nenhum serviço associado foi encontrado no Link Inicial do fluxo.")
    if not dados["campos"]:
        logger.warning("Nenhum campo de preenchimento foi encontrado no fluxo.")
    if not dados["scripts"]:
        logger.warning("Nenhum script (IronPython) ou expressão de Gateway foi encontrado no fluxo.")

    return dados


# ==========================================================
# COMPARAÇÃO ENTRE DUAS VERSÕES DO MESMO FLUXO (v2)
# ==========================================================
def extrair_versao_do_nome_arquivo(caminho_xml):
    """Devolve o número da versão presente no nome do arquivo ('..._Versão_45_...') ou None."""
    nome_base = os.path.splitext(os.path.basename(caminho_xml))[0]
    nome_normalizado = re.sub(r"\s+", " ", nome_base.replace("_", " "))
    match = re.search(r"Vers[aã]o\s+(\d+)", nome_normalizado, flags=re.IGNORECASE)
    return match.group(1) if match else None


def _indexar(itens, chave):
    """Indexa uma lista de dicionários por uma chave, preservando a ordem de entrada."""
    indice = {}
    for item in itens:
        indice[chave(item)] = item
    return indice


def _comparar_colecao(antes, depois, chave, campos_comparados):
    """
    Compara duas coleções de dicionários e classifica cada item em
    incluído (só na nova), removido (só na antiga) ou modificado (existe nas duas,
    mas algum dos `campos_comparados` mudou).

    Um item MODIFICADO carrega 'antes'/'depois' e a lista de campos que diferem, para
    o documento poder dizer exatamente o que mudou.
    """
    idx_antes = _indexar(antes, chave)
    idx_depois = _indexar(depois, chave)

    incluidos = [item for k, item in idx_depois.items() if k not in idx_antes]
    removidos = [item for k, item in idx_antes.items() if k not in idx_depois]

    modificados = []
    for k, item_depois in idx_depois.items():
        item_antes = idx_antes.get(k)
        if item_antes is None:
            continue
        diferencas = [c for c in campos_comparados if item_antes.get(c) != item_depois.get(c)]
        if diferencas:
            modificados.append({
                "chave": k,
                "antes": item_antes,
                "depois": item_depois,
                "diferencas": diferencas,
            })

    return {
        "incluidos": incluidos,
        "removidos": removidos,
        "modificados": modificados,
        # Estado completo da versão NOVA, para as seções que devem ser listadas por
        # inteiro no artefato (não só o delta) — caso dos Serviços.
        "todos": list(idx_depois.values()),
        "chaves_incluidas": {chave(item) for item in incluidos},
    }


def comparar_dados(dados_antes, dados_depois, logger=None):
    """
    Confronta o resultado de extrair_dados_xml() de duas versões do MESMO fluxo e
    devolve apenas o delta entre elas, por seção.

    Serviços e itens de configuração (anexos) não têm conteúdo além da própria
    identidade, então só podem ser incluídos ou removidos — nunca 'modificados'.
    """
    comparacao = {
        "nome_fluxo": dados_depois.get("nome_fluxo") or dados_antes.get("nome_fluxo", ""),
        "nome_fluxo_antes": dados_antes.get("nome_fluxo", ""),
        "nome_fluxo_mudou": dados_antes.get("nome_fluxo", "") != dados_depois.get("nome_fluxo", ""),
        "servicos": _comparar_colecao(
            dados_antes["servicos"], dados_depois["servicos"],
            chave=lambda s: (s["tipo"], s["nome"]),
            campos_comparados=[],
        ),
        # ATENÇÃO: 'rotulo' está deliberadamente FORA da comparação. Ele vem do <Rotulo>
        # da tela e o mesmo campo pode ter rótulos diferentes (inclusive vazio) em telas
        # distintas do mesmo XML; como a extração guarda o primeiro que encontra, comparar
        # rótulo acusaria "alteração" só porque a ordem de varredura mudou. Os demais campos
        # vêm do <CustomProperty> (definição global do campo), que é estável.
        "campos": _comparar_colecao(
            dados_antes["campos"], dados_depois["campos"],
            chave=lambda c: c["nome"],
            campos_comparados=["descricao", "tipo", "tabela", "controle", "lista_itens"],
        ),
        "anexos": _comparar_colecao(
            dados_antes["anexos"], dados_depois["anexos"],
            chave=lambda a: a["sigla"],
            campos_comparados=[],
        ),
        "scripts": _comparar_colecao(
            dados_antes["scripts"], dados_depois["scripts"],
            chave=lambda s: s["local"],
            campos_comparados=["codigo"],
        ),
    }

    if logger:
        if comparacao["nome_fluxo_mudou"]:
            logger.warning(
                f"O nome do subprocesso difere entre os dois XMLs: "
                f"'{comparacao['nome_fluxo_antes']}' -> '{comparacao['nome_fluxo']}'. "
                f"Confirme que os arquivos são do MESMO fluxo."
            )
        for secao in ("servicos", "campos", "anexos", "scripts"):
            r = comparacao[secao]
            logger.info(
                f"Comparação de {secao}: {len(r['incluidos'])} incluído(s), "
                f"{len(r['removidos'])} removido(s), {len(r['modificados'])} modificado(s)."
            )

    return comparacao


def descrever_mudanca_lista(item_antes, item_depois):
    """
    Monta o texto da coluna 'Lista de Itens': as opções da versão atual e, quando a
    lista mudou, quais opções entraram e quais saíram.
    """
    itens_depois = separar_lista_itens((item_depois or {}).get("lista_itens", ""))
    itens_antes = separar_lista_itens((item_antes or {}).get("lista_itens", ""))

    if not itens_depois and not itens_antes:
        return ""

    partes = ["; ".join(itens_depois) if itens_depois else "(sem itens)"]

    incluidos = [i for i in itens_depois if i not in itens_antes]
    removidos = [i for i in itens_antes if i not in itens_depois]
    if incluidos:
        partes.append("Incluídos: " + "; ".join(incluidos))
    if removidos:
        partes.append("Removidos: " + "; ".join(removidos))

    return "\n".join(partes)


def descrever_tipo_campo(campo):
    """'String' vira 'String (DropDownList)' quando o campo é de seleção/lista."""
    tipo = (campo or {}).get("tipo", "") or "String"
    controle = (campo or {}).get("controle", "")
    if controle and controle in ("DropDownList", "ListBox", "CheckBox", "SearchList", "RadioButton"):
        return f"{tipo} ({controle})"
    return tipo


def descrever_mudanca_tipo(item_antes, item_depois):
    """
    Texto da coluna 'Tipo' na tabela de Campos alterados. Quando o tipo mudou entre as
    versões (ex.: o campo era String e virou Integer), mostra a transição
    'String -> Integer' em vez de apenas o tipo atual, para a alteração ficar explícita
    no artefato. Se o tipo não mudou, mostra só o tipo atual.
    """
    tipo_antes = descrever_tipo_campo(item_antes)
    tipo_depois = descrever_tipo_campo(item_depois)
    if item_antes is not None and tipo_antes != tipo_depois:
        return f"{tipo_antes} → {tipo_depois}"
    return tipo_depois


class SupravizioDocApp:

    TOTAL_ETAPAS = 6  # nº de estágios que _set_status percorre durante gerar_documento

    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Artefato Supravizio")
        self.root.geometry("700x800")
        self.root.resizable(True, True)

        self.xml_path = ""
        self.template_path = ""
        self.output_dir = ""

        # Aba de comparação (v2): os dois XMLs do mesmo fluxo em versões diferentes.
        self.xml_antes_path = ""
        self.xml_depois_path = ""

        self.logger, self.log_handler = setup_logger()

        # Cobre exceções fora do try/except de gerar_documento: erros em callbacks de
        # botão/diálogo durante o mainloop() (o Tk não propaga isso para sys.excepthook
        # por padrão) e qualquer coisa que escape do __init__ antes da GUI existir.
        manipulador_excecoes = criar_manipulador_excecoes(self.logger, self.log_handler)
        self.root.report_callback_exception = manipulador_excecoes

        # Guarda o que foi preenchido automaticamente, para poder substituir esses valores
        # ao trocar de XML sem sobrescrever o que o usuário digitou à mão.
        self.macro_auto = ""
        self.proc_auto = ""
        self.macro_auto_cmp = ""
        self.proc_auto_cmp = ""

        # =========================
        # ABAS
        # =========================
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill="both", expand=True)

        aba_artefato = tk.Frame(self.notebook)
        self.notebook.add(aba_artefato, text="  Artefato de Fluxo  ")
        self._construir_aba_artefato(aba_artefato)

        aba_comparacao = tk.Frame(self.notebook)
        self.notebook.add(aba_comparacao, text="  Comparar Versões  ")
        self._construir_aba_comparacao(aba_comparacao)

        self._carregar_historico()

    # ==========================================================
    # ABA 1 — ARTEFATO DE UM ÚNICO FLUXO (v1)
    # ==========================================================
    def _construir_aba_artefato(self, root):
        tk.Label(root, text="Automação de Artefatos - Supravizio", font=("Arial", 14, "bold")).pack(pady=10)

        frame_inputs = tk.Frame(root)
        frame_inputs.pack(pady=5, fill="x", padx=20)

        tk.Label(frame_inputs, text="Macroprocesso:", anchor="w").pack(fill="x")
        self.entry_macro = tk.Entry(frame_inputs)
        self.entry_macro.pack(pady=2, fill="x")

        tk.Label(frame_inputs, text="Processo:", anchor="w").pack(fill="x")
        self.entry_proc = tk.Entry(frame_inputs)
        self.entry_proc.pack(pady=2, fill="x")

        self.lbl_deteccao = tk.Label(
            frame_inputs, text="", fg="gray", anchor="w", font=("Arial", 8, "italic"), wraplength=580, justify="left"
        )
        self.lbl_deteccao.pack(fill="x", pady=(0, 4))

        tk.Label(frame_inputs, text="Descrição da Alteração/Criação:", anchor="w").pack(fill="x")
        self.text_descricao = tk.Text(frame_inputs, height=4, width=50)
        self.text_descricao.pack(pady=2, fill="x")

        tk.Label(frame_inputs, text="Evidências em Homologação (Nº Chamado):", anchor="w").pack(fill="x")
        self.entry_evidencia = tk.Entry(frame_inputs)
        self.entry_evidencia.pack(pady=2, fill="x")

        # =========================
        # SELETORES DE DIRETÓRIOS E ARQUIVOS
        # =========================
        frame_files = tk.Frame(root)
        frame_files.pack(pady=10, fill="x", padx=20)

        texto_btn_xml = "1. Selecionar XML do Fluxo"
        if DND_DISPONIVEL:
            texto_btn_xml += "  (ou arraste o arquivo aqui)"
        self.btn_xml = tk.Button(frame_files, text=texto_btn_xml, command=self.load_xml)
        self.btn_xml.pack(fill="x", pady=2)
        self.lbl_xml = tk.Label(frame_files, text="Nenhum arquivo XML selecionado", fg="gray", anchor="w")
        self.lbl_xml.pack(fill="x", pady=(0, 10))
        self._lbl_xml_bg_padrao = self.lbl_xml.cget("bg")

        if DND_DISPONIVEL:
            self._registrar_alvo_drop((self.btn_xml, self.lbl_xml), self.lbl_xml, self._definir_xml)

        self.btn_template = tk.Button(frame_files, text="2. Selecionar Template DOCX", command=self.load_template)
        self.btn_template.pack(fill="x", pady=2)
        self.lbl_template = tk.Label(frame_files, text="Nenhum template DOCX selecionado", fg="gray", anchor="w")
        self.lbl_template.pack(fill="x", pady=(0, 10))

        self.btn_dir = tk.Button(frame_files, text="3. Selecionar Pasta de Destino", command=self.load_dir)
        self.btn_dir.pack(fill="x", pady=2)
        self.lbl_dir = tk.Label(frame_files, text="Nenhuma pasta selecionada", fg="gray", anchor="w")
        self.lbl_dir.pack(fill="x")

        self.btn_gerar = tk.Button(
            root, text="Gerar Artefato DOCX", font=("Arial", 12, "bold"),
            bg="#4CAF50", fg="white", command=self.gerar_documento
        )
        self.btn_gerar.pack(pady=15, fill="x", padx=50)

        # =========================
        # STATUS / PROGRESSO
        # =========================
        self.lbl_status = tk.Label(root, text="", fg="gray", anchor="w")
        self.lbl_status.pack(fill="x", padx=50)

        self.progress = ttk.Progressbar(root, mode="determinate", maximum=self.TOTAL_ETAPAS)
        self.progress.pack(fill="x", padx=50, pady=(0, 10))

    # ==========================================================
    # ABA 2 — COMPARAÇÃO ENTRE DUAS VERSÕES DO MESMO FLUXO (v2)
    # ==========================================================
    def _construir_aba_comparacao(self, root):
        tk.Label(
            root, text="Comparar Versões do Mesmo Fluxo", font=("Arial", 14, "bold")
        ).pack(pady=(10, 2))
        tk.Label(
            root,
            text="Gera o artefato contendo APENAS o que mudou entre as duas versões.",
            fg="gray", font=("Arial", 9)
        ).pack(pady=(0, 8))

        frame_inputs = tk.Frame(root)
        frame_inputs.pack(pady=5, fill="x", padx=20)

        tk.Label(frame_inputs, text="Macroprocesso:", anchor="w").pack(fill="x")
        self.entry_macro_cmp = tk.Entry(frame_inputs)
        self.entry_macro_cmp.pack(pady=2, fill="x")

        tk.Label(frame_inputs, text="Processo:", anchor="w").pack(fill="x")
        self.entry_proc_cmp = tk.Entry(frame_inputs)
        self.entry_proc_cmp.pack(pady=2, fill="x")

        self.lbl_deteccao_cmp = tk.Label(
            frame_inputs, text="", fg="gray", anchor="w", font=("Arial", 8, "italic"),
            wraplength=620, justify="left"
        )
        self.lbl_deteccao_cmp.pack(fill="x", pady=(0, 4))

        tk.Label(frame_inputs, text="Descrição da Alteração:", anchor="w").pack(fill="x")
        self.text_descricao_cmp = tk.Text(frame_inputs, height=3, width=50)
        self.text_descricao_cmp.pack(pady=2, fill="x")

        tk.Label(frame_inputs, text="Evidências em Homologação (Nº Chamado):", anchor="w").pack(fill="x")
        self.entry_evidencia_cmp = tk.Entry(frame_inputs)
        self.entry_evidencia_cmp.pack(pady=2, fill="x")

        frame_files = tk.Frame(root)
        frame_files.pack(pady=10, fill="x", padx=20)

        sufixo_dnd = "  (ou arraste aqui)" if DND_DISPONIVEL else ""

        self.btn_xml_antes = tk.Button(
            frame_files, text=f"1. XML da versão ANTERIOR{sufixo_dnd}", command=self.load_xml_antes
        )
        self.btn_xml_antes.pack(fill="x", pady=2)
        self.lbl_xml_antes = tk.Label(frame_files, text="Nenhum XML selecionado", fg="gray", anchor="w")
        self.lbl_xml_antes.pack(fill="x", pady=(0, 8))

        self.btn_xml_depois = tk.Button(
            frame_files, text=f"2. XML da versão NOVA{sufixo_dnd}", command=self.load_xml_depois
        )
        self.btn_xml_depois.pack(fill="x", pady=2)
        self.lbl_xml_depois = tk.Label(frame_files, text="Nenhum XML selecionado", fg="gray", anchor="w")
        self.lbl_xml_depois.pack(fill="x", pady=(0, 8))

        if DND_DISPONIVEL:
            self._registrar_alvo_drop(
                (self.btn_xml_antes, self.lbl_xml_antes), self.lbl_xml_antes, self._definir_xml_antes
            )
            self._registrar_alvo_drop(
                (self.btn_xml_depois, self.lbl_xml_depois), self.lbl_xml_depois, self._definir_xml_depois
            )

        self.btn_template_cmp = tk.Button(
            frame_files, text="3. Selecionar Template DOCX", command=self.load_template
        )
        self.btn_template_cmp.pack(fill="x", pady=2)
        self.lbl_template_cmp = tk.Label(
            frame_files, text="Nenhum template DOCX selecionado", fg="gray", anchor="w"
        )
        self.lbl_template_cmp.pack(fill="x", pady=(0, 8))

        self.btn_dir_cmp = tk.Button(
            frame_files, text="4. Selecionar Pasta de Destino", command=self.load_dir
        )
        self.btn_dir_cmp.pack(fill="x", pady=2)
        self.lbl_dir_cmp = tk.Label(frame_files, text="Nenhuma pasta selecionada", fg="gray", anchor="w")
        self.lbl_dir_cmp.pack(fill="x")

        self.btn_gerar_cmp = tk.Button(
            root, text="Gerar Artefato Comparativo", font=("Arial", 12, "bold"),
            bg="#1565C0", fg="white", command=self.gerar_documento_comparacao
        )
        self.btn_gerar_cmp.pack(pady=15, fill="x", padx=50)

        self.lbl_status_cmp = tk.Label(root, text="", fg="gray", anchor="w")
        self.lbl_status_cmp.pack(fill="x", padx=50)

        self.progress_cmp = ttk.Progressbar(root, mode="determinate", maximum=self.TOTAL_ETAPAS)
        self.progress_cmp.pack(fill="x", padx=50, pady=(0, 10))

    # ==========================================================
    # HISTÓRICO (settings.json)
    # ==========================================================
    def _carregar_historico(self):
        """Pré-preenche a interface com o que foi usado na última execução."""
        self.config = carregar_config(self.logger)

        template_salvo = self.config.get("template_path", "")
        if template_salvo and os.path.isfile(template_salvo):
            self.template_path = template_salvo

        dir_salvo = self.config.get("output_dir", "")
        if dir_salvo and os.path.isdir(dir_salvo):
            self.output_dir = dir_salvo

        self._refrescar_labels_compartilhados()

        # Passa pelo MESMO mecanismo de preenchimento inteligente que a detecção via XML
        # usa: só entra se o campo estiver vazio (primeira abertura do app), nunca
        # sobrescreve algo que o usuário já tenha digitado nesta sessão.
        macro_salvo = self.config.get("macroprocesso", "")
        proc_salvo = self.config.get("processo", "")
        self.macro_auto = self.aplicar_valor_automatico(
            self.entry_macro, macro_salvo, self.macro_auto, "Macroprocesso"
        )
        self.proc_auto = self.aplicar_valor_automatico(
            self.entry_proc, proc_salvo, self.proc_auto, "Processo"
        )
        self.macro_auto_cmp = self.aplicar_valor_automatico(
            self.entry_macro_cmp, macro_salvo, self.macro_auto_cmp, "Macroprocesso (comparação)"
        )
        self.proc_auto_cmp = self.aplicar_valor_automatico(
            self.entry_proc_cmp, proc_salvo, self.proc_auto_cmp, "Processo (comparação)"
        )

    def _refrescar_labels_compartilhados(self):
        """
        Template e pasta de destino são os mesmos para as duas abas; este método mantém
        os rótulos das duas sincronizados quando um deles é alterado.
        """
        if self.template_path:
            texto_template, cor = os.path.basename(self.template_path), "black"
        else:
            texto_template, cor = "Nenhum template DOCX selecionado", "gray"
        self.lbl_template.config(text=texto_template, fg=cor)
        self.lbl_template_cmp.config(text=texto_template, fg=cor)

        if self.output_dir:
            texto_dir, cor_dir = self.output_dir, "black"
        else:
            texto_dir, cor_dir = "Nenhuma pasta selecionada", "gray"
        self.lbl_dir.config(text=texto_dir, fg=cor_dir)
        self.lbl_dir_cmp.config(text=texto_dir, fg=cor_dir)

    def _atualizar_config(self, **campos):
        """Atualiza e persiste campos do histórico (nunca lança em caso de falha)."""
        self.config.update(campos)
        salvar_config(self.config, self.logger)

    # ==========================================================
    # LOG
    # ==========================================================
    def gravar_log(self, nome_fluxo=""):
        """Grava um arquivo de log para a geração recém-concluída e esvazia o buffer."""
        linhas = self.log_handler.drenar()
        return escrever_arquivo_log(linhas, nome_fluxo if nome_fluxo else "Supravizio")

    # ==========================================================
    # FUNÇÕES DE INTERFACE
    # ==========================================================
    def load_xml(self):
        caminho = filedialog.askopenfilename(
            title="Selecionar XML",
            filetypes=[("XML", "*.xml")],
            initialdir=self.config.get("last_xml_dir") or None
        )
        if caminho:
            self._definir_xml(caminho)

    def _definir_xml(self, caminho):
        """
        Único ponto que efetivamente adota um caminho de XML como self.xml_path —
        usado tanto pelo diálogo de seleção (load_xml) quanto pelo arrastar-e-soltar
        (_ao_soltar_xml), para as duas formas passarem pela MESMA detecção automática.
        """
        if not caminho.lower().endswith(".xml"):
            self.logger.warning(f"Arquivo descartado por não ser .xml: {caminho}")
            messagebox.showwarning("Arquivo inválido", "Selecione (ou arraste) um arquivo .xml.")
            return
        if not os.path.isfile(caminho):
            self.logger.warning(f"Caminho de XML não é um arquivo existente: {caminho}")
            messagebox.showwarning("Arquivo inválido", f"Arquivo não encontrado:\n{caminho}")
            return
        try:
            self.xml_path = caminho
            self.lbl_xml.config(text=os.path.basename(self.xml_path), fg="black")
            self.logger.info(f"XML selecionado: {self.xml_path}")
            self._atualizar_config(last_xml_dir=os.path.dirname(self.xml_path))
            self.tentar_preencher_macro_processo()
        except Exception as e:
            self.logger.exception(f"Falha ao processar o XML selecionado '{caminho}': {e}")
            messagebox.showerror("Erro ao ler XML", f"Não foi possível ler o arquivo selecionado:\n{e}")

    # ----- Aba de comparação: XML anterior e XML novo -----
    def _validar_xml(self, caminho):
        """Valida um caminho de XML antes de adotá-lo; devolve True/False e avisa o usuário."""
        if not caminho.lower().endswith(".xml"):
            self.logger.warning(f"Arquivo descartado por não ser .xml: {caminho}")
            messagebox.showwarning("Arquivo inválido", "Selecione (ou arraste) um arquivo .xml.")
            return False
        if not os.path.isfile(caminho):
            self.logger.warning(f"Caminho de XML não é um arquivo existente: {caminho}")
            messagebox.showwarning("Arquivo inválido", f"Arquivo não encontrado:\n{caminho}")
            return False
        return True

    def load_xml_antes(self):
        caminho = filedialog.askopenfilename(
            title="Selecionar XML da versão ANTERIOR",
            filetypes=[("XML", "*.xml")],
            initialdir=self.config.get("last_xml_dir") or None
        )
        if caminho:
            self._definir_xml_antes(caminho)

    def load_xml_depois(self):
        caminho = filedialog.askopenfilename(
            title="Selecionar XML da versão NOVA",
            filetypes=[("XML", "*.xml")],
            initialdir=self.config.get("last_xml_dir") or None
        )
        if caminho:
            self._definir_xml_depois(caminho)

    def _definir_xml_antes(self, caminho):
        if not self._validar_xml(caminho):
            return
        self.xml_antes_path = caminho
        self.lbl_xml_antes.config(text=self._rotulo_versao(caminho), fg="black")
        self.logger.info(f"XML da versão anterior selecionado: {caminho}")
        self._atualizar_config(last_xml_dir=os.path.dirname(caminho))

    def _definir_xml_depois(self, caminho):
        if not self._validar_xml(caminho):
            return
        self.xml_depois_path = caminho
        self.lbl_xml_depois.config(text=self._rotulo_versao(caminho), fg="black")
        self.logger.info(f"XML da versão nova selecionado: {caminho}")
        self._atualizar_config(last_xml_dir=os.path.dirname(caminho))
        # O Macroprocesso/Processo do artefato são os da versão NOVA.
        self.tentar_preencher_macro_processo_cmp()

    def _rotulo_versao(self, caminho):
        """Mostra o nome do arquivo e, quando dá para detectar, a versão ('Versão 45')."""
        versao = extrair_versao_do_nome_arquivo(caminho)
        nome = os.path.basename(caminho)
        return f"[Versão {versao}] {nome}" if versao else nome

    # ==========================================================
    # ARRASTAR E SOLTAR (tkinterdnd2, se disponível)
    # ==========================================================
    def _registrar_alvo_drop(self, widgets, label_destaque, ao_definir):
        """
        Torna `widgets` uma zona de soltar arquivos: destaca `label_destaque` durante o
        arraste e entrega o caminho solto a `ao_definir` — o mesmo método que o botão de
        seleção usa, para não haver dois caminhos de código divergentes.
        """
        bg_padrao = label_destaque.cget("bg")

        def ao_entrar(event):
            label_destaque.config(bg="#E3F2FD")

        def ao_sair(event):
            label_destaque.config(bg=bg_padrao)

        def ao_soltar(event):
            label_destaque.config(bg=bg_padrao)
            caminhos = dividir_caminhos_dnd(event.data)
            if caminhos:
                ao_definir(caminhos[0])

        for widget in widgets:
            widget.drop_target_register(DND_FILES)
            widget.dnd_bind("<<DropEnter>>", ao_entrar)
            widget.dnd_bind("<<DropLeave>>", ao_sair)
            widget.dnd_bind("<<Drop>>", ao_soltar)

    # ==========================================================
    # DETECÇÃO AUTOMÁTICA DE PROCESSO / SUBPROCESSO
    # ==========================================================
    def extrair_processo_do_nome_arquivo(self, caminho_xml):
        return extrair_processo_do_nome_arquivo(caminho_xml)

    def extrair_macroprocesso_do_xml(self, root):
        return extrair_macroprocesso_do_xml(root)

    def aplicar_valor_automatico(self, entry, novo_valor, valor_auto_anterior, rotulo):
        """
        Atualiza um campo preenchido automaticamente ao trocar de XML.
        Substitui o valor quando o campo está vazio ou ainda contém o que foi detectado
        para o XML anterior; se o usuário digitou algo à mão, o texto dele é preservado.
        Retorna o novo valor automático a ser lembrado.
        """
        atual = entry.get().strip()
        if atual and atual != valor_auto_anterior:
            self.logger.info(
                f"Campo '{rotulo}' foi editado manualmente ('{atual}'); "
                f"o valor detectado ('{novo_valor}') não foi aplicado."
            )
            return valor_auto_anterior

        entry.delete(0, tk.END)
        if novo_valor:
            entry.insert(0, novo_valor)
        return novo_valor

    def tentar_preencher_macro_processo(self):
        processo, sub_do_nome = self.extrair_processo_do_nome_arquivo(self.xml_path)

        macro = None
        try:
            macro = self.extrair_macroprocesso_do_xml(self.ler_xml_root())
        except Exception as e:
            self.logger.warning(f"Não foi possível ler o XML para buscar o Macroprocesso: {e}")

        avisos = []

        if processo:
            self.logger.info(f"Processo detectado pelo nome do arquivo: '{processo}'")
            self.proc_auto = self.aplicar_valor_automatico(self.entry_proc, processo, self.proc_auto, "Processo")
        else:
            self.logger.warning(
                f"Não foi possível detectar o Processo pelo nome do arquivo "
                f"'{os.path.basename(self.xml_path)}' (padrão esperado: "
                f"'Processo_Versão_N_Subprocesso'). Preencha manualmente."
            )
            self.proc_auto = self.aplicar_valor_automatico(self.entry_proc, "", self.proc_auto, "Processo")
            avisos.append("Processo")

        if sub_do_nome:
            self.logger.info(f"Subprocesso indicado pelo nome do arquivo: '{sub_do_nome}'")

        if macro:
            self.logger.info(f"Macroprocesso encontrado no XML: '{macro}'")
            self.macro_auto = self.aplicar_valor_automatico(self.entry_macro, macro, self.macro_auto, "Macroprocesso")
        else:
            self.logger.warning(
                "Macroprocesso não consta no XML exportado pelo Supravizio nem no nome do "
                "arquivo. Preencha manualmente."
            )
            self.macro_auto = self.aplicar_valor_automatico(self.entry_macro, "", self.macro_auto, "Macroprocesso")
            avisos.append("Macroprocesso")

        if avisos:
            self.lbl_deteccao.config(
                text="Preencha manualmente: " + ", ".join(avisos)
                     + ". O Subprocesso é lido direto do XML.",
                fg="#B71C1C"
            )
        else:
            self.lbl_deteccao.config(
                text="Macroprocesso e Processo preenchidos automaticamente (confira antes de gerar).",
                fg="#2E7D32"
            )

    def tentar_preencher_macro_processo_cmp(self):
        """Mesma detecção da aba 1, porém baseada no XML da versão NOVA."""
        caminho = self.xml_depois_path
        processo, sub_do_nome = extrair_processo_do_nome_arquivo(caminho)

        macro = None
        try:
            macro = extrair_macroprocesso_do_xml(ler_xml_root(caminho))
        except Exception as e:
            self.logger.warning(f"Não foi possível ler o XML para buscar o Macroprocesso: {e}")

        avisos = []

        if processo:
            self.logger.info(f"Processo detectado pelo nome do arquivo (comparação): '{processo}'")
            self.proc_auto_cmp = self.aplicar_valor_automatico(
                self.entry_proc_cmp, processo, self.proc_auto_cmp, "Processo (comparação)"
            )
        else:
            self.logger.warning(
                f"Não foi possível detectar o Processo pelo nome do arquivo "
                f"'{os.path.basename(caminho)}'. Preencha manualmente."
            )
            self.proc_auto_cmp = self.aplicar_valor_automatico(
                self.entry_proc_cmp, "", self.proc_auto_cmp, "Processo (comparação)"
            )
            avisos.append("Processo")

        if sub_do_nome:
            self.logger.info(f"Subprocesso indicado pelo nome do arquivo (comparação): '{sub_do_nome}'")

        if macro:
            self.logger.info(f"Macroprocesso encontrado no XML (comparação): '{macro}'")
            self.macro_auto_cmp = self.aplicar_valor_automatico(
                self.entry_macro_cmp, macro, self.macro_auto_cmp, "Macroprocesso (comparação)"
            )
        else:
            self.logger.warning(
                "Macroprocesso não consta no XML nem no nome do arquivo. Preencha manualmente."
            )
            self.macro_auto_cmp = self.aplicar_valor_automatico(
                self.entry_macro_cmp, "", self.macro_auto_cmp, "Macroprocesso (comparação)"
            )
            avisos.append("Macroprocesso")

        if avisos:
            self.lbl_deteccao_cmp.config(
                text="Preencha manualmente: " + ", ".join(avisos)
                     + ". O Subprocesso é lido direto do XML.",
                fg="#B71C1C"
            )
        else:
            self.lbl_deteccao_cmp.config(
                text="Macroprocesso e Processo preenchidos automaticamente (confira antes de gerar).",
                fg="#2E7D32"
            )

    def load_template(self):
        caminho = filedialog.askopenfilename(title="Selecionar DOCX", filetypes=[("Word", "*.docx")])
        if not caminho:
            return
        try:
            self.template_path = caminho
            self._refrescar_labels_compartilhados()
            self.logger.info(f"Template selecionado: {self.template_path}")
            self._atualizar_config(template_path=self.template_path)
        except Exception as e:
            self.logger.exception(f"Falha ao selecionar o template '{caminho}': {e}")
            messagebox.showerror("Erro ao selecionar template", f"Não foi possível usar o arquivo selecionado:\n{e}")

    def load_dir(self):
        caminho = filedialog.askdirectory(title="Selecionar Destino")
        if not caminho:
            return
        try:
            self.output_dir = caminho
            self._refrescar_labels_compartilhados()
            self.logger.info(f"Pasta de destino selecionada: {self.output_dir}")
            self._atualizar_config(output_dir=self.output_dir)
        except Exception as e:
            self.logger.exception(f"Falha ao selecionar a pasta de destino '{caminho}': {e}")
            messagebox.showerror("Erro ao selecionar pasta", f"Não foi possível usar a pasta selecionada:\n{e}")

    # ==========================================================
    # LÓGICA DE EXTRAÇÃO DE DADOS
    # ==========================================================
    def get_texto(self, node, tag, default=""):
        return get_texto(node, tag, default)

    def extrair_propriedades_campos(self, root):
        return extrair_propriedades_campos(root)

    def ler_xml_root(self):
        return ler_xml_root(self.xml_path)

    def extrair_dados_xml(self):
        return extrair_dados_xml(self.xml_path, self.logger)

    def add_code_block(self, paragraph, code_text):
        """Cria um bloco de código com aparência de IDE, incluindo sintaxe colorida."""

        # ------------------------------------------------------------------
        # Aparência do bloco
        # ------------------------------------------------------------------
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.left_indent = Cm(0.15)
        paragraph.paragraph_format.right_indent = Cm(0.15)

        pPr = paragraph._p.get_or_add_pPr()

        # Fundo cinza nativo XML
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F4F4F4') 
        pPr.append(shd)

        # Borda discreta ao redor do código
        pBdr = OxmlElement('w:pBdr')
        for side in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '4')
            border.set(qn('w:color'), 'CCCCCC')
            pBdr.append(border)
        pPr.append(pBdr)

        # ------------------------------------------------------------------
        # Highlight de sintaxe
        # ------------------------------------------------------------------
        cores = {
            Token.Keyword: "0000FF",
            Token.Keyword.Constant: "0000FF",
            Token.Keyword.Namespace: "0000FF",
            Token.Name.Builtin: "267F99",
            Token.Name.Function: "795E26",
            Token.Name.Class: "267F99",
            Token.Name.Decorator: "795E26",
            Token.String: "A31515",
            Token.String.Doc: "A31515",
            Token.Number: "098658",
            Token.Comment: "008000",
            Token.Operator: "000000",
            Token.Punctuation: "000000",
            Token.Name: "000000",
            Token.Text: "000000",
        }

        linhas_do_codigo = code_text.split('\n')
        
        for k, linha_bruta in enumerate(linhas_do_codigo):
            if not linha_bruta.strip():
                if k < len(linhas_do_codigo) - 1:
                    paragraph.add_run().add_break()
                continue
                
            espacos_iniciais = len(linha_bruta) - len(linha_bruta.lstrip(' '))
            if espacos_iniciais > 0:
                 run_espaco = paragraph.add_run(" " * espacos_iniciais)
                 run_espaco.font.name = "Consolas"
                 run_espaco.font.size = Pt(9)
                 
            for token_type, token_text in lex(linha_bruta.lstrip(' '), PythonLexer()):
                cor = "000000"
                tipo = token_type
                while tipo is not Token:
                    if tipo in cores:
                        cor = cores[tipo]
                        break
                    tipo = tipo.parent

                run = paragraph.add_run(token_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string(cor)

                # Força compatibilidade da fonte Consolas
                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:ascii"), "Consolas")
                rFonts.set(qn("w:hAnsi"), "Consolas")
                rFonts.set(qn("w:cs"), "Consolas")

            if k < len(linhas_do_codigo) - 1:
                paragraph.add_run().add_break()

    def inserir_paragrafo_depois(self, paragraph):
        """Cria um parágrafo vazio logo após `paragraph`, herdando sua formatação."""
        novo_p = OxmlElement("w:p")
        pPr_origem = paragraph._p.find(qn("w:pPr"))
        if pPr_origem is not None:
            novo_p.append(copy.deepcopy(pPr_origem))
        paragraph._p.addnext(novo_p)
        return Paragraph(novo_p, paragraph._parent)

    def localizar_placeholder(self, doc, idx_titulo):
        """
        Devolve o índice do parágrafo vazio que o template deixa logo abaixo de um título,
        para que o conteúdo seja escrito NELE em vez de numa linha nova depois dele
        (o que deixava uma linha em branco entre o título e o texto).
        """
        if idx_titulo == -1:
            return -1
        idx = idx_titulo + 1
        if idx < len(doc.paragraphs) and not doc.paragraphs[idx].text.strip():
            # O placeholder pode conter runs vazios com quebras de linha, que empurrariam
            # o texto para baixo do título; remove-os mantendo a formatação do parágrafo.
            for run in list(doc.paragraphs[idx].runs):
                run._r.getparent().remove(run._r)
            return idx
        return -1

    def _atualizar_progresso(self, lbl, progress, texto, passo):
        """
        Atualiza o texto de status e a barra de progresso, e força o Tk a repintar
        AGORA (sem isso a janela só atualizaria no fim da geração, já que tudo roda de
        forma síncrona na thread principal — não há threading aqui).
        """
        lbl.config(text=texto)
        progress["value"] = passo
        self.root.update_idletasks()

    def _set_status(self, texto, passo):
        """Progresso da aba 1 (artefato de um único fluxo)."""
        self._atualizar_progresso(self.lbl_status, self.progress, texto, passo)

    def _set_status_cmp(self, texto, passo):
        """Progresso da aba 2 (comparação entre versões)."""
        self._atualizar_progresso(self.lbl_status_cmp, self.progress_cmp, texto, passo)

    def gerar_documento(self):
        if not self.xml_path or not self.template_path or not self.output_dir:
            messagebox.showerror("Erro", "Preencha todos os caminhos (XML, Template e Pasta).")
            return

        self.logger.info("=== Iniciando geração do artefato ===")
        self.logger.info(f"XML: {self.xml_path}")
        self.logger.info(f"Template: {self.template_path}")
        self.logger.info(f"Pasta de destino: {self.output_dir}")

        nome_fluxo = ""
        try:
            self._set_status("Lendo e extraindo dados do XML...", 1)
            dados = self.extrair_dados_xml()
            nome_fluxo = dados["nome_fluxo"]

            macro = self.entry_macro.get().strip()
            proc = self.entry_proc.get().strip()
            desc = self.text_descricao.get("1.0", tk.END).strip()
            evid = self.entry_evidencia.get().strip()

            if not macro:
                self.logger.warning("Campo 'Macroprocesso' não foi preenchido (nem manual, nem automaticamente).")
            if not proc:
                self.logger.warning("Campo 'Processo' não foi preenchido (nem manual, nem automaticamente).")
            if not evid:
                self.logger.warning("Campo 'Evidências em Homologação' não foi preenchido.")

            self._set_status("Carregando template DOCX...", 2)
            doc = Document(self.template_path)

            self._set_status("Preenchendo campos do documento...", 3)
            idx_titulo_desc = -1
            idx_titulo_servico = -1
            idx_evid = -1

            for i, p in enumerate(doc.paragraphs):
                texto = p.text.strip()
                if texto.startswith("Macroprocesso:"):
                    p.text = ""
                    p.add_run("Macroprocesso: ").bold = True
                    p.add_run(macro)
                elif texto.startswith("Processo:"):
                    p.text = ""
                    p.add_run("Processo: ").bold = True
                    p.add_run(proc)
                elif texto.startswith("Sub-processos:"):
                    p.text = ""
                    p.add_run("Sub-processos: ").bold = True
                    p.add_run(dados["nome_fluxo"] if dados["nome_fluxo"] else "")
                elif texto == "Descrição":
                    idx_titulo_desc = i
                elif texto == "Serviço":
                    idx_titulo_servico = i
                elif texto.startswith("Evidências em Homologação:"):
                    idx_evid = i
                    p.text = ""
                    p.add_run("Evidências em Homologação: ").bold = True
                    p.add_run(evid)

            idx_desc = self.localizar_placeholder(doc, idx_titulo_desc)
            idx_servico = self.localizar_placeholder(doc, idx_titulo_servico)

            if idx_desc == -1:
                self.logger.warning("Parágrafo em branco após o título 'Descrição' não encontrado; descrição não inserida.")
            if idx_servico == -1:
                self.logger.warning("Parágrafo em branco após o título 'Serviço' não encontrado; serviços não inseridos.")
            if idx_evid == -1:
                self.logger.warning("Âncora 'Evidências em Homologação:' não encontrada no template DOCX; scripts não inseridos.")
            if len(doc.tables) == 0:
                self.logger.warning("Template DOCX não possui tabelas; tabela de Campos não preenchida.")
            if len(doc.tables) <= 2:
                self.logger.warning("Template DOCX não possui a 3ª tabela esperada; tabela de Anexos não preenchida.")

            self._set_status("Formatando blocos de script...", 4)
            if idx_evid != -1 and dados["scripts"]:
                estilo_ancora = doc.paragraphs[idx_evid].style
                for s_dict in reversed(dados["scripts"]):
                    doc.paragraphs[idx_evid].insert_paragraph_before("", style=estilo_ancora)
                    
                    p_code = doc.paragraphs[idx_evid].insert_paragraph_before("")
                    self.add_code_block(p_code, s_dict["codigo"])
                    
                    p_local = doc.paragraphs[idx_evid].insert_paragraph_before("", style=estilo_ancora)
                    p_local.add_run(s_dict["local"]).bold = True

            # Escreve no parágrafo-placeholder do template (preservando o recuo dele) e só
            # cria linhas extras para os itens seguintes, para o texto colar logo abaixo do título.
            if idx_servico != -1 and dados["servicos"]:
                p_atual = doc.paragraphs[idx_servico]
                p_atual.add_run(f"- Tipo: {dados['servicos'][0]['tipo']} | Serviço: {dados['servicos'][0]['nome']}")
                for srv in dados["servicos"][1:]:
                    p_atual = self.inserir_paragrafo_depois(p_atual)
                    p_atual.add_run(f"- Tipo: {srv['tipo']} | Serviço: {srv['nome']}")

            if idx_desc != -1 and desc:
                linhas_desc = desc.split("\n")
                p_atual = doc.paragraphs[idx_desc]
                p_atual.add_run(linhas_desc[0])
                for linha in linhas_desc[1:]:
                    p_atual = self.inserir_paragrafo_depois(p_atual)
                    p_atual.add_run(linha)

            self._set_status("Montando tabelas...", 5)
            if len(doc.tables) > 0:
                t_campos = doc.tables[0]
                for row in t_campos.rows[1:]: t_campos._element.remove(row._tr)
                for c in dados["campos"]:
                    row = t_campos.add_row()
                    row.cells[0].text = c["nome"]
                    row.cells[1].text = c.get("descricao", c["rotulo"])
                    row.cells[2].text = c.get("tipo", "String")
                    row.cells[3].text = c["tabela"]
                    row.cells[4].text = "Sim"

            if len(doc.tables) > 2:
                t_anexos = doc.tables[2]
                for row in t_anexos.rows[1:]: t_anexos._element.remove(row._tr)
                for a in dados["anexos"]:
                    row = t_anexos.add_row()
                    row.cells[0].text = a["nome"]
                    row.cells[1].text = a["sigla"]
                    row.cells[2].text = "Sim"

            nome_arquivo = limpar_nome_arquivo(dados["nome_fluxo"])
            caminho_completo = os.path.join(self.output_dir, f"Artefato_{nome_arquivo}.docx")

            self._set_status("Salvando documento...", 6)
            doc.save(caminho_completo)
            self.logger.info(f"Artefato salvo com sucesso em: {caminho_completo}")
            self.logger.info("=== Geração concluída ===")

            # Só persiste Macroprocesso/Processo depois de uma geração bem-sucedida
            # (não a cada tecla), para não gravar valores que nunca chegaram a virar artefato.
            self._atualizar_config(macroprocesso=macro, processo=proc)

            log_path = self.gravar_log(nome_fluxo)
            self._set_status("Concluído.", 0)
            messagebox.showinfo(
                "Sucesso",
                f"Artefato Injetado com formatação visual de código!\n\nSalvo em:\n{caminho_completo}"
                f"\n\nLog desta execução:\n{log_path}"
            )

        except Exception as e:
            self.logger.exception(f"Falha na geração do artefato: {e}")
            log_path = self.gravar_log(nome_fluxo)
            self._set_status("Falha na geração.", 0)
            messagebox.showerror(
                "Erro",
                f"Falha na manipulação do DOCX:\n{e}\n\nDetalhes em:\n{log_path}"
            )

    # ==========================================================
    # GERAÇÃO DO ARTEFATO COMPARATIVO (aba 2)
    # ==========================================================
    def gerar_documento_comparacao(self):
        if not self.xml_antes_path or not self.xml_depois_path:
            messagebox.showerror("Erro", "Selecione os dois XMLs (versão anterior e versão nova).")
            return
        if not self.template_path or not self.output_dir:
            messagebox.showerror("Erro", "Selecione o Template DOCX e a Pasta de destino.")
            return
        if os.path.abspath(self.xml_antes_path) == os.path.abspath(self.xml_depois_path):
            messagebox.showerror("Erro", "Os dois XMLs são o mesmo arquivo. Selecione versões diferentes.")
            return

        self.logger.info("=== Iniciando geração do artefato COMPARATIVO ===")
        self.logger.info(f"XML anterior: {self.xml_antes_path}")
        self.logger.info(f"XML novo: {self.xml_depois_path}")
        self.logger.info(f"Template: {self.template_path}")
        self.logger.info(f"Pasta de destino: {self.output_dir}")

        nome_fluxo = ""
        try:
            self._set_status_cmp("Lendo XML da versão anterior...", 1)
            dados_antes = extrair_dados_xml(self.xml_antes_path, self.logger)

            self._set_status_cmp("Lendo XML da versão nova...", 2)
            dados_depois = extrair_dados_xml(self.xml_depois_path, self.logger)

            self._set_status_cmp("Comparando as duas versões...", 3)
            comp = comparar_dados(dados_antes, dados_depois, self.logger)
            nome_fluxo = comp["nome_fluxo"]

            if comp["nome_fluxo_mudou"]:
                seguir = messagebox.askyesno(
                    "Fluxos diferentes?",
                    f"O nome do subprocesso é diferente nos dois arquivos:\n\n"
                    f"Anterior: {comp['nome_fluxo_antes']}\n"
                    f"Novo: {comp['nome_fluxo']}\n\n"
                    f"Eles podem não ser o mesmo fluxo. Deseja continuar mesmo assim?"
                )
                if not seguir:
                    self.logger.info("Geração cancelada pelo usuário (nomes de subprocesso divergentes).")
                    self._set_status_cmp("Cancelado.", 0)
                    self.gravar_log(nome_fluxo)
                    return

            macro = self.entry_macro_cmp.get().strip()
            proc = self.entry_proc_cmp.get().strip()
            desc = self.text_descricao_cmp.get("1.0", tk.END).strip()
            evid = self.entry_evidencia_cmp.get().strip()

            if not macro:
                self.logger.warning("Campo 'Macroprocesso' não foi preenchido.")
            if not proc:
                self.logger.warning("Campo 'Processo' não foi preenchido.")
            if not evid:
                self.logger.warning("Campo 'Evidências em Homologação' não foi preenchido.")

            self._set_status_cmp("Carregando template DOCX...", 4)
            doc = Document(self.template_path)

            self._set_status_cmp("Preenchendo o que mudou...", 5)
            self._preencher_documento_comparacao(doc, comp, macro, proc, desc, evid)

            versao_antes = extrair_versao_do_nome_arquivo(self.xml_antes_path)
            versao_depois = extrair_versao_do_nome_arquivo(self.xml_depois_path)
            nome_arquivo = limpar_nome_arquivo(nome_fluxo)
            if versao_antes and versao_depois:
                nome_arquivo = f"{nome_arquivo}_v{versao_antes}_para_v{versao_depois}"
            else:
                nome_arquivo = f"{nome_arquivo}_Comparacao"
            caminho_completo = os.path.join(self.output_dir, f"Artefato_{nome_arquivo}.docx")

            self._set_status_cmp("Salvando documento...", 6)
            doc.save(caminho_completo)
            self.logger.info(f"Artefato comparativo salvo com sucesso em: {caminho_completo}")
            self.logger.info("=== Geração concluída ===")

            self._atualizar_config(macroprocesso=macro, processo=proc)

            resumo = " | ".join(
                f"{secao}: +{len(comp[secao]['incluidos'])} "
                f"-{len(comp[secao]['removidos'])} ~{len(comp[secao]['modificados'])}"
                for secao in ("servicos", "campos", "anexos", "scripts")
            )
            log_path = self.gravar_log(nome_fluxo)
            self._set_status_cmp("Concluído.", 0)
            messagebox.showinfo(
                "Sucesso",
                f"Artefato comparativo gerado!\n\nSalvo em:\n{caminho_completo}"
                f"\n\nMudanças encontradas (incluídos/removidos/modificados):\n{resumo}"
                f"\n\nLog desta execução:\n{log_path}"
            )

        except Exception as e:
            self.logger.exception(f"Falha na geração do artefato comparativo: {e}")
            log_path = self.gravar_log(nome_fluxo)
            self._set_status_cmp("Falha na geração.", 0)
            messagebox.showerror(
                "Erro",
                f"Falha ao gerar o artefato comparativo:\n{e}\n\nDetalhes em:\n{log_path}"
            )

    def _preencher_documento_comparacao(self, doc, comp, macro, proc, desc, evid):
        """
        Escreve no template apenas o delta entre as duas versões.

        Divisão das tabelas do modelo:
          - "Campos do processo" (tabela 0): campos que entraram (Sim) ou saíram (Não);
          - "Campos alterados"  (tabela 1): campos que existem nas duas versões mas
            mudaram (tipicamente a lista de um combobox) — sempre 'Sim', pois continuam
            existindo na versão atual;
          - "Itens de Configuração" (tabela 2): anexos que entraram ou saíram.
        """
        idx_titulo_desc = -1
        idx_titulo_servico = -1
        idx_evid = -1

        for i, p in enumerate(doc.paragraphs):
            texto = p.text.strip()
            if texto.startswith("Macroprocesso:"):
                p.text = ""
                p.add_run("Macroprocesso: ").bold = True
                p.add_run(macro)
            elif texto.startswith("Processo:"):
                p.text = ""
                p.add_run("Processo: ").bold = True
                p.add_run(proc)
            elif texto.startswith("Sub-processos:"):
                p.text = ""
                p.add_run("Sub-processos: ").bold = True
                p.add_run(comp["nome_fluxo"])
            elif texto == "Descrição":
                idx_titulo_desc = i
            elif texto == "Serviço":
                idx_titulo_servico = i
            elif texto.startswith("Evidências em Homologação:"):
                idx_evid = i
                p.text = ""
                p.add_run("Evidências em Homologação: ").bold = True
                p.add_run(evid)

        idx_desc = self.localizar_placeholder(doc, idx_titulo_desc)
        idx_servico = self.localizar_placeholder(doc, idx_titulo_servico)

        if idx_evid == -1:
            self.logger.warning("Âncora 'Evidências em Homologação:' não encontrada; scripts não inseridos.")
        if len(doc.tables) < 3:
            self.logger.warning(
                f"Template tem apenas {len(doc.tables)} tabela(s); esperadas 3 "
                f"(Campos do processo, Campos alterados, Itens de Configuração)."
            )

        # ----- SCRIPTS (incluídos, removidos e alterados) -----
        scripts_doc = []
        for s in comp["scripts"]["incluidos"]:
            scripts_doc.append((f"[INCLUÍDO] {s['local']}", s["codigo"]))
        for m in comp["scripts"]["modificados"]:
            scripts_doc.append((f"[ALTERADO] {m['chave']}", m["depois"]["codigo"]))
        for s in comp["scripts"]["removidos"]:
            scripts_doc.append((f"[REMOVIDO] {s['local']}", s["codigo"]))

        if idx_evid != -1 and scripts_doc:
            estilo_ancora = doc.paragraphs[idx_evid].style
            for local, codigo in reversed(scripts_doc):
                doc.paragraphs[idx_evid].insert_paragraph_before("", style=estilo_ancora)
                p_code = doc.paragraphs[idx_evid].insert_paragraph_before("")
                self.add_code_block(p_code, codigo)
                p_local = doc.paragraphs[idx_evid].insert_paragraph_before("", style=estilo_ancora)
                p_local.add_run(local).bold = True

        # ----- SERVIÇOS -----
        # Diferente das demais seções, os serviços são listados POR INTEIRO (todos os da
        # versão nova), e não apenas o delta: o artefato precisa mostrar quais serviços o
        # fluxo atende hoje. Os que entraram/saíram ficam marcados.
        linhas_servico = []
        incluidos_srv = comp["servicos"]["chaves_incluidas"]
        for s in comp["servicos"]["todos"]:
            marca = "[INCLUÍDO] " if (s["tipo"], s["nome"]) in incluidos_srv else ""
            linhas_servico.append(f"- {marca}Tipo: {s['tipo']} | Serviço: {s['nome']}")
        for s in comp["servicos"]["removidos"]:
            linhas_servico.append(f"- [REMOVIDO] Tipo: {s['tipo']} | Serviço: {s['nome']}")
        if not linhas_servico:
            linhas_servico.append("- Nenhum serviço associado ao fluxo.")

        if idx_servico != -1:
            p_atual = doc.paragraphs[idx_servico]
            p_atual.add_run(linhas_servico[0])
            for linha in linhas_servico[1:]:
                p_atual = self.inserir_paragrafo_depois(p_atual)
                p_atual.add_run(linha)
        else:
            self.logger.warning("Parágrafo em branco após 'Serviço' não encontrado; serviços não inseridos.")

        # ----- DESCRIÇÃO -----
        if idx_desc != -1 and desc:
            linhas_desc = desc.split("\n")
            p_atual = doc.paragraphs[idx_desc]
            p_atual.add_run(linhas_desc[0])
            for linha in linhas_desc[1:]:
                p_atual = self.inserir_paragrafo_depois(p_atual)
                p_atual.add_run(linha)
        elif idx_desc == -1:
            self.logger.warning("Parágrafo em branco após 'Descrição' não encontrado; descrição não inserida.")

        # ----- TABELA 0: Campos do processo (incluídos / removidos) -----
        if len(doc.tables) > 0:
            t_campos = doc.tables[0]
            for row in t_campos.rows[1:]:
                t_campos._element.remove(row._tr)
            for c in comp["campos"]["incluidos"]:
                self._escrever_linha(t_campos, [
                    c["nome"], c.get("descricao", ""), descrever_tipo_campo(c), c.get("tabela", ""), "Sim"
                ])
            for c in comp["campos"]["removidos"]:
                self._escrever_linha(t_campos, [
                    c["nome"], c.get("descricao", ""), descrever_tipo_campo(c), c.get("tabela", ""), "Não"
                ])
            if not comp["campos"]["incluidos"] and not comp["campos"]["removidos"]:
                self._escrever_linha(t_campos, ["N/A", "Nenhum campo incluído ou removido", "N/A", "N/A", "N/A"])

        # ----- TABELA 1: Campos alterados (existem nas duas, mudaram) -----
        if len(doc.tables) > 1:
            t_alterados = doc.tables[1]
            for row in t_alterados.rows[1:]:
                t_alterados._element.remove(row._tr)
            for m in comp["campos"]["modificados"]:
                depois = m["depois"]
                self._escrever_linha(t_alterados, [
                    depois["nome"],
                    descrever_mudanca_tipo(m["antes"], depois),
                    descrever_mudanca_lista(m["antes"], depois),
                    "Sim",
                ])
                self.logger.info(
                    f"Campo alterado '{depois['nome']}': mudou {', '.join(m['diferencas'])}."
                )
            if not comp["campos"]["modificados"]:
                self._escrever_linha(t_alterados, ["N/A", "N/A", "Nenhum campo alterado", "N/A"])

        # ----- TABELA 2: Itens de Configuração (anexos) -----
        if len(doc.tables) > 2:
            t_anexos = doc.tables[2]
            for row in t_anexos.rows[1:]:
                t_anexos._element.remove(row._tr)
            for a in comp["anexos"]["incluidos"]:
                self._escrever_linha(t_anexos, [a["nome"], a["sigla"], "Sim"])
            for a in comp["anexos"]["removidos"]:
                self._escrever_linha(t_anexos, [a["nome"], a["sigla"], "Não"])
            if not comp["anexos"]["incluidos"] and not comp["anexos"]["removidos"]:
                self._escrever_linha(t_anexos, ["N/A", "Nenhum item alterado", "N/A"])

    def _escrever_linha(self, tabela, valores):
        """Acrescenta uma linha à tabela, respeitando o nº de colunas do modelo."""
        row = tabela.add_row()
        for i, valor in enumerate(valores):
            if i < len(row.cells):
                row.cells[i].text = valor
        return row


if __name__ == "__main__":
    def _excecao_nao_tratada_inicial(exc_type, exc_value, exc_tb):
        """
        Cobre falhas fora do alcance de report_callback_exception (ex.: durante a
        construção de SupravizioDocApp, antes do mainloop existir), quando ainda não
        há self.logger/self.log_handler para reaproveitar.
        """
        import traceback
        linhas = [linha.rstrip("\n") for linha in traceback.format_exception(exc_type, exc_value, exc_tb)]
        log_path = escrever_arquivo_log(linhas, "crash")
        try:
            messagebox.showerror(
                "Erro inesperado",
                f"Ocorreu um erro inesperado e o programa será encerrado:\n{exc_value}"
                f"\n\nDetalhes em:\n{log_path}"
            )
        except Exception:
            pass

    sys.excepthook = _excecao_nao_tratada_inicial

    # TkinterDnD.Tk() é um substituto direto de tk.Tk() que só acrescenta suporte a
    # arrastar-e-soltar; sem a lib instalada, cai para o tk.Tk() normal (sem DnD).
    root = TkinterDnD.Tk() if DND_DISPONIVEL else tk.Tk()
    app = SupravizioDocApp(root)
    root.mainloop()
